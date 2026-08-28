import asyncio
import logging
import math
import os
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional
import aiofiles
from fastapi import UploadFile, HTTPException, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.config import settings
from app.db.base import AsyncSessionLocal
from app.db.models.document import Document, DocumentChunk

logger = logging.getLogger("datapilot.document_service")

ALLOWED_DOCUMENT_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".json"}


def _tokenize(text: str) -> List[str]:
    """Tokenize and normalize text into clean words."""
    words = re.findall(r"\b[a-zA-Z0-9_]{2,}\b", text.lower())
    return words


def _compute_tf_vector(tokens: List[str]) -> Dict[str, float]:
    """Calculate normalized term-frequency vector for cosine similarity."""
    tf: Dict[str, float] = {}
    for t in tokens:
        tf[t] = tf.get(t, 0.0) + 1.0
    
    # Normalize vector to unit length
    magnitude = math.sqrt(sum(v * v for v in tf.values()))
    if magnitude > 0:
        for k in tf:
            tf[k] = round(tf[k] / magnitude, 4)
    return tf


def _cosine_similarity(vec1: Dict[str, float], vec2: Dict[str, float]) -> float:
    """Calculate cosine similarity between two term-frequency vector dictionaries."""
    common_keys = set(vec1.keys()) & set(vec2.keys())
    if not common_keys:
        return 0.0
    return sum(vec1[k] * vec2[k] for k in common_keys)


def _sanitize_text(text: str) -> str:
    """Sanitizes text by removing control characters, unicode replacement chars, and unreadable binary artifacts."""
    if not text:
        return ""
    # Remove null bytes and non-printable control characters (except newline, tab, carriage return)
    cleaned = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F\uFFFD]", " ", text)
    # Collapse multiple whitespace characters
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _is_legible_text(text: str, min_alpha_ratio: float = 0.5) -> bool:
    """Checks if extracted text contains a reasonable ratio of readable alphanumeric characters."""
    if not text or len(text.strip()) < 10:
        return False
    alpha_chars = sum(1 for c in text if c.isalnum() or c.isspace() or c in ".,!?:;\"'()-$%&/")
    return (alpha_chars / len(text)) >= min_alpha_ratio


def _extract_text_from_file(file_path: str, extension: str) -> str:
    """Extract raw text from PDF, TXT, Markdown, or other supported formats."""
    ext = extension.lower()
    path = Path(file_path)

    if not path.exists():
        return ""

    if ext in (".txt", ".md", ".json"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return _sanitize_text(f.read())
    
    elif ext == ".pdf":
        text_content = []
        try:
            # Attempt to use pypdf if available
            import pypdf
            reader = pypdf.PdfReader(str(path))
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted and _is_legible_text(extracted):
                    text_content.append(_sanitize_text(extracted))
        except Exception:
            # Fallback binary stream text scraper with strict legibility check
            try:
                with open(path, "rb") as f:
                    raw = f.read().decode("latin1", errors="ignore")
                    matches = re.findall(r"\(([^\)]+)\)\s*Tj", raw)
                    if matches:
                        cand = " ".join(matches)
                        if _is_legible_text(cand):
                            text_content.append(_sanitize_text(cand))
            except Exception:
                pass
        return _sanitize_text("\n\n".join(text_content)) if text_content else ""

    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(str(path))
            return _sanitize_text("\n".join([p.text for p in doc.paragraphs if p.text and _is_legible_text(p.text)]))
        except Exception:
            return ""

    return ""


def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 60) -> List[str]:
    """Split text into overlapping semantic passages."""
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks = []
    current_chunk = ""

    for p in paragraphs:
        if len(current_chunk) + len(p) <= chunk_size:
            current_chunk = f"{current_chunk}\n\n{p}".strip()
        else:
            if current_chunk:
                chunks.append(current_chunk)
            # If paragraph itself is longer than chunk_size, split by characters
            if len(p) > chunk_size:
                start = 0
                while start < len(p):
                    chunks.append(p[start:start + chunk_size])
                    start += chunk_size - overlap
                current_chunk = ""
            else:
                current_chunk = p

    if current_chunk:
        chunks.append(current_chunk)

    return chunks if chunks else [text[:chunk_size]]


async def save_document_file(
    file: UploadFile,
    workspace_id: str,
    user_id: str,
    db: AsyncSession,
) -> Document:
    """Save an uploaded business document file and create the database Document record."""
    filename = file.filename or "uploaded_document"
    ext = os.path.splitext(filename)[1].lower()

    if ext not in ALLOWED_DOCUMENT_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported document format '{ext}'. Allowed: {', '.join(ALLOWED_DOCUMENT_EXTENSIONS)}",
        )

    # Prepare storage directory
    upload_dir = Path(settings.upload_dir) / "documents"
    upload_dir.mkdir(parents=True, exist_ok=True)

    unique_name = f"{uuid.uuid4()}{ext}"
    saved_path = upload_dir / unique_name

    # Save to disk asynchronously
    file_size = 0
    async with aiofiles.open(saved_path, "wb") as out_file:
        while content := await file.read(1024 * 1024):  # 1MB chunks
            file_size += len(content)
            await out_file.write(content)

    # Create Document record
    title = os.path.splitext(filename)[0].replace("_", " ").replace("-", " ").title()
    document = Document(
        workspace_id=workspace_id,
        uploaded_by=user_id,
        title=title,
        original_filename=filename,
        file_path=str(saved_path),
        file_size_bytes=file_size,
        mime_type=file.content_type or "application/octet-stream",
        file_extension=ext,
        chunk_count=0,
        status="UPLOADED",
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)
    return document


async def process_document_background(document_id: str):
    """Background worker to extract text, create chunks, and generate search embeddings."""
    async with AsyncSessionLocal() as db:
        try:
            result = await db.execute(select(Document).where(Document.id == document_id))
            document = result.scalar_one_or_none()
            if not document:
                logger.error(f"Document {document_id} not found for background chunking.")
                return

            document.status = "PROCESSING"
            await db.commit()

            # Extract raw text
            text = _extract_text_from_file(document.file_path, document.file_extension)
            chunks = _chunk_text(text)

            # Create chunk entities with vector embeddings
            for idx, chunk_text in enumerate(chunks):
                tokens = _tokenize(chunk_text)
                tf_vec = _compute_tf_vector(tokens)

                chunk_obj = DocumentChunk(
                    document_id=document.id,
                    chunk_index=idx + 1,
                    content=chunk_text,
                    token_count=len(tokens),
                    chunk_metadata={"length": len(chunk_text), "page": (idx // 3) + 1},
                    embedding=list(tf_vec.items()),  # Store as key-value pairs
                )
                db.add(chunk_obj)

            document.chunk_count = len(chunks)
            document.status = "INDEXED"
            await db.commit()
            logger.info(f"Document {document.title} indexed successfully with {len(chunks)} chunks.")

        except Exception as e:
            logger.exception(f"Document processing failed for {document_id}: {e}")
            try:
                result = await db.execute(select(Document).where(Document.id == document_id))
                document = result.scalar_one_or_none()
                if document:
                    document.status = "ERROR"
                    document.error_message = str(e)
                    await db.commit()
            except Exception:
                pass


async def search_workspace_documents(
    workspace_id: str,
    query: str,
    limit: int = 5,
    db: AsyncSession = None,
) -> List[Dict[str, Any]]:
    """Hybrid semantic vector search across all indexed documents in a workspace."""
    query_tokens = _tokenize(query)
    query_vec = _compute_tf_vector(query_tokens)

    # Fetch all chunks belonging to documents in the workspace
    stmt = (
        select(DocumentChunk, Document)
        .join(Document, Document.id == DocumentChunk.document_id)
        .where(
            Document.workspace_id == workspace_id,
            Document.status == "INDEXED",
            Document.is_deleted == False,
        )
    )
    result = await db.execute(stmt)
    rows = result.all()

    scored_results = []
    for chunk, doc in rows:
        cleaned_content = _sanitize_text(chunk.content)
        if not _is_legible_text(cleaned_content):
            continue

        # Reconstruct vector
        chunk_vec = dict(chunk.embedding) if chunk.embedding else {}
        similarity = _cosine_similarity(query_vec, chunk_vec)
        
        # Word overlap boost
        overlap_count = len(set(query_tokens) & set(chunk_vec.keys()))
        final_score = similarity * 0.7 + (overlap_count / max(len(query_tokens), 1)) * 0.3

        if final_score >= 0.10:
            scored_results.append({
                "chunk_id": chunk.id,
                "document_id": doc.id,
                "document_title": doc.title,
                "content": cleaned_content,
                "similarity_score": round(final_score, 4),
                "chunk_index": chunk.chunk_index,
                "chunk_metadata": chunk.chunk_metadata,
            })

    # Sort descending by score
    scored_results.sort(key=lambda x: x["similarity_score"], reverse=True)
    return scored_results[:limit]
